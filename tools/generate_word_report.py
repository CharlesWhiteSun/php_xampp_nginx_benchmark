#!/usr/bin/env python3
"""
Generate a Word (.docx) management-friendly deployment recommendation report
based on an existing HTML benchmark report (reports/report_*.html).

Usage:
  python tools/generate_word_report.py [path/to/report.html]

If no path is provided, the script uses the latest report_*.html under reports/.

Requires: python-docx, beautifulsoup4
"""
import sys
from pathlib import Path
import re

try:
    from bs4 import BeautifulSoup
    from docx import Document
    from docx.shared import Pt
except Exception as e:
    print("Missing Python dependencies: bs4 or python-docx.\nInstall with: pip install python-docx beautifulsoup4", file=sys.stderr)
    raise


def find_latest_report(reports_dir: Path) -> Path:
    files = sorted(reports_dir.glob('report_*.html'), key=lambda p: p.stat().st_mtime, reverse=True)
    if not files:
        raise FileNotFoundError('No report_*.html found in reports/')
    return files[0]


def extract_text_from_html(html_path: Path) -> dict:
    html = html_path.read_text(encoding='utf-8')
    soup = BeautifulSoup(html, 'html.parser')

    data = {}
    # Title
    h1 = soup.find('h1')
    data['title'] = h1.get_text(strip=True) if h1 else 'Benchmark Report'

    # Meta: generated at / source
    meta_area = html
    m = re.search(r'generated_at["\']?:\s*"([0-9\- :]+)"', html)
    if m:
        data['generated_at'] = m.group(1)
    else:
        data['generated_at'] = None

    # Try to find summary section content: look for headings containing 'Summary' or '摘要'
    summary_text = ''
    for header in soup.find_all(['h2','h3','h4']):
        if re.search(r'Summary|摘要|Configuration Summary|Configuration', header.get_text(), re.I):
            # collect sibling paragraphs
            sibs = []
            for sib in header.find_next_siblings():
                if sib.name and sib.name.startswith('h'):
                    break
                sibs.append(sib.get_text(separator=' ', strip=True))
            summary_text = '\n'.join([s for s in sibs if s.strip()])
            if summary_text:
                break

    data['summary'] = summary_text or 'No explicit summary section found in HTML. See source report.'

    # Extract Peak Client Concurrency if present in text
    peak = None
    m2 = re.search(r'Peak Client Concurrency[^\d]*(\d{1,6})', html)
    if m2:
        peak = int(m2.group(1))
    else:
        # fallback: find numbers near 'Peak'
        m3 = re.search(r'Peak[^\d]{0,60}(\d{3,6})', html)
        if m3:
            peak = int(m3.group(1))

    data['peak_concurrency'] = peak

    # Collect notable metrics: Request/sec and p99 from any table or text
    # Grep for 'Requests per second' lines
    reqs = re.findall(r'Requests per second:\s*([0-9]+\.?[0-9]*)', html)
    data['requests_per_sec_samples'] = reqs

    p99s = re.findall(r'99%\s*[<>=]*\s*(\d+)', html)
    data['p99_samples'] = p99s

    return data


def make_recommendation(data: dict) -> str:
    peak = data.get('peak_concurrency')
    lines = []
    lines.append('Recommendation overview:')
    if peak is not None:
        lines.append(f'- Observed peak client concurrency (approx): {peak}')

    # Simple decision logic
    if peak is None:
        lines.append('- Unable to detect a clear peak concurrency value from the source report; recommend conservative, horizontally scalable deployment (Kubernetes) for production.)')
        rec = 'Kubernetes (managed)'
    elif peak < 200:
        lines.append('- Low concurrency observed -> Docker Compose or single-node Docker + nginx is sufficient for this scale. Consider using Docker Compose or Nomad for simple orchestration.')
        rec = 'Docker Compose / Nomad (single-node)'
    elif peak < 1000:
        lines.append('- Medium concurrency -> use an orchestrator with autoscaling and service discovery. Consider lightweight Kubernetes (k3s) or managed Kubernetes (EKS/GKE/AKS).')
        rec = 'Kubernetes (k3s or managed)'
    else:
        lines.append('- High concurrency (>=1000) -> production-grade orchestration required with horizontal autoscaling, robust networking, and observability. Recommend managed Kubernetes (EKS/GKE/AKS) with HPA + Cluster Autoscaler, or HashiCorp Nomad with autoscaler for specific constraints.')
        rec = 'Managed Kubernetes (EKS/GKE/AKS)'

    lines.append(f'Primary recommendation: {rec}')

    # Add specific software choices
    lines.append('\nSuggested stack:')
    lines.append('- Reverse proxy / load balancer: NGINX or Traefik (NGINX for stability/perf; Traefik for dynamic config)')
    lines.append('- Container runtime: Docker')
    lines.append('- Orchestration: See above (Compose for dev, k3s/k8s for prod)')
    lines.append('- PHP process manager: php-fpm with tuned pm.* settings and separated pools for high concurrency')
    lines.append('- Observability: Prometheus + Grafana, and centralized logs (ELK/EFK)')

    # Operational notes
    lines.append('\nOperational considerations:')
    lines.append('- Ensure connection limits (nginx worker_connections, php-fpm pm.max_children) align with expected concurrency')
    lines.append('- Use readiness/liveness probes and graceful shutdown for zero-downtime deploys')
    lines.append('- Use health-check endpoints and circuit breakers when integrating external I/O')

    return '\n'.join(lines)


def build_docx(data: dict, recommendation: str, html_path: Path, out_path: Path) -> None:
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Noto Sans'
    font.size = Pt(11)

    doc.add_heading('Deployment Recommendation Report', level=1)
    doc.add_paragraph(f'Source report: {html_path}')
    if data.get('generated_at'):
        doc.add_paragraph(f'Report generated at: {data.get("generated_at")}')

    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph(data.get('summary', ''))

    doc.add_heading('Key Metrics Extracted', level=2)
    if data.get('peak_concurrency'):
        doc.add_paragraph(f'Peak Client Concurrency: {data.get("peak_concurrency")}')
    if data.get('requests_per_sec_samples'):
        doc.add_paragraph('Requests per second samples: ' + ', '.join(data.get('requests_per_sec_samples')))
    if data.get('p99_samples'):
        doc.add_paragraph('p99 latency samples (ms): ' + ', '.join(data.get('p99_samples')))

    doc.add_heading('Recommendation', level=2)
    for line in recommendation.splitlines():
        doc.add_paragraph(line)

    doc.add_heading('Actionable Next Steps', level=2)
    steps = [
        '1) Proof-of-concept: Deploy with Docker Compose + nginx-multi + php-fpm-multi on a staging node; validate with existing benchmark scripts.',
        '2) For production: provision k3s or managed Kubernetes, deploy via Helm charts, enable HPA, configure Ingress with NGINX Ingress Controller.',
        '3) Configure monitoring (Prometheus + Grafana) and alerts for latency and error-rate spikes.',
        '4) Tune php-fpm and nginx settings according to observed concurrency and resources.'
    ]
    for s in steps:
        doc.add_paragraph(s)

    doc.add_page_break()
    doc.save(str(out_path))


def main():
    reports_dir = Path('reports')
    if len(sys.argv) > 1:
        html_path = Path(sys.argv[1])
    else:
        html_path = find_latest_report(reports_dir)

    data = extract_text_from_html(html_path)
    recommendation = make_recommendation(data)

    out_name = f"deployment_recommendation_{html_path.stem}.docx"
    out_path = Path('reports') / out_name
    build_docx(data, recommendation, html_path, out_path)
    print(f'Generated Word report: {out_path}')


if __name__ == '__main__':
    main()
